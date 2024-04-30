import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
import shutil

def clone_paragraph(original_paragraph, modified_doc):
    new_paragraph = modified_doc.add_paragraph()
    for run in original_paragraph.runs:
        new_run = new_paragraph.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.size = run.font.size
        new_run.font.name = run.font.name
        new_run.font.color.rgb = run.font.color.rgb
    new_paragraph.alignment = original_paragraph.alignment
    return new_paragraph

def modify_cover_letter(position_name, company_name):
    # # Original document path
    # input_file_path = r"C:\Users\njain\OneDrive - Cal State Fullerton\SPRING 2024\Nilay Jain Resume\Cover Letter\Cover_Letter.docx"
    
    # create a copy of the input file
    shutil.copyfile("Cover_Letter.docx", "Modified_Cover_Letter.docx")

    # # Load the original document
    # original_doc = Document(input_file_path)

    # Create a new document for modification
    modified_doc = Document('Modified_Cover_Letter.docx')

    # Clone content from original document to modified document
    for paragraph in modified_doc.paragraphs:
        # new_paragraph = clone_paragraph(paragraph, modified_doc)
        if '{Position_Name}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{Position_Name}', position_name)
        if '{Company_Name}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{Company_Name}', company_name)
    
    # adjusting the font size.
    if modified_doc.paragraphs:
        first_paragraph = modified_doc.paragraphs[5]
        for run in first_paragraph.runs:
            run.font.size = Pt(13)  # Adjust the font size as needed

    # Adjust font size to fit on a single page
    for paragraph in modified_doc.paragraphs:
        paragraph.space_after = Pt(0.5)  # Adjust the spacing between paragraphs as needed

    # Save the modified document with a different name
    modified_doc.save('Modified_Cover_Letter.docx')

    # Convert the modified document to PDF
    output_pdf_path = "CoverLetter.pdf"
    convert('Modified_Cover_Letter.docx', output_pdf_path)

if __name__ == "__main__":
    # Check if there are enough command-line arguments
    if len(sys.argv) < 3:
        print("Usage: python script.py [position_name] [company_name]")
        sys.exit(1)

    # Get the position name and company name from command-line arguments
    position_name = sys.argv[1]
    company_name = sys.argv[2]
    modify_cover_letter(position_name, company_name)
